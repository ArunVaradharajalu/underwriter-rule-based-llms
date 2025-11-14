#
#    Copyright 2024 IBM Corp.
#
#    Licensed under the Apache License, Version 2.0 (the "License");
#    you may not use this file except in compliance with the License.
#    You may obtain a copy of the License at
#
#        http://www.apache.org/licenses/LICENSE-2.0
#
#    Unless required by applicable law or agreed to in writing, software
#    distributed under the License is distributed on an "AS IS" BASIS,
#    WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#    See the License for the specific language governing permissions and
#    limitations under the License.
#
import os
import tempfile
from typing import Dict, Optional
import boto3
from io import BytesIO

class DocumentExtractor:
    """
    Multi-format document text extraction service
    Supports: PDF, Excel (.xlsx, .xls), Word (.docx), and text files
    """

    def __init__(self):
        """Initialize AWS S3 client for document retrieval"""
        self.aws_access_key = os.getenv("AWS_ACCESS_KEY_ID")
        self.aws_secret_key = os.getenv("AWS_SECRET_ACCESS_KEY")
        self.aws_region = os.getenv("AWS_REGION", "us-east-1")

        self.isConfigured = self.aws_access_key is not None and self.aws_secret_key is not None

        if self.isConfigured:
            try:
                self.s3_client = boto3.client(
                    's3',
                    aws_access_key_id=self.aws_access_key,
                    aws_secret_access_key=self.aws_secret_key,
                    region_name=self.aws_region
                )
                print(f"DocumentExtractor initialized with S3 access")
            except Exception as e:
                print(f"Error initializing S3 client: {e}")
                self.isConfigured = False
        else:
            print("DocumentExtractor: AWS credentials not configured")
            self.s3_client = None

    def extract_text_from_s3(self, s3_url: str) -> Dict[str, str]:
        """
        Extract text from a document stored in S3
        Automatically detects format and uses appropriate extractor

        Args:
            s3_url: S3 URL in format s3://bucket/key or https://bucket.s3.region.amazonaws.com/key

        Returns:
            Dictionary with:
                - text: Extracted text content
                - format: Detected file format (pdf, excel, word, text)
                - error: Error message if extraction failed
        """
        if not self.isConfigured:
            return {"error": "AWS S3 not configured", "text": "", "format": "unknown"}

        try:
            # Parse S3 URL
            if s3_url.startswith('s3://'):
                # Format: s3://bucket/key
                parts = s3_url.replace('s3://', '').split('/', 1)
                s3_bucket = parts[0]
                s3_key = parts[1] if len(parts) > 1 else ""
            elif 's3.amazonaws.com' in s3_url or 's3.' in s3_url:
                # Format: https://bucket.s3.region.amazonaws.com/key
                parts = s3_url.split('/', 3)
                s3_bucket = parts[2].split('.')[0]
                s3_key = parts[3] if len(parts) > 3 else ""
            else:
                return {"error": f"Invalid S3 URL format: {s3_url}", "text": "", "format": "unknown"}

            # Detect file format from extension
            file_extension = os.path.splitext(s3_key)[1].lower()

            print(f"Extracting text from S3: {s3_bucket}/{s3_key} (format: {file_extension})")

            # Download file from S3
            response = self.s3_client.get_object(Bucket=s3_bucket, Key=s3_key)
            file_content = response['Body'].read()

            # Route to appropriate extractor
            if file_extension == '.pdf':
                text = self._extract_from_pdf(file_content)
                return {"text": text, "format": "pdf", "s3_bucket": s3_bucket, "s3_key": s3_key}

            elif file_extension in ['.xlsx', '.xls']:
                text = self._extract_from_excel(file_content)
                return {"text": text, "format": "excel", "s3_bucket": s3_bucket, "s3_key": s3_key}

            elif file_extension == '.docx':
                text = self._extract_from_word(file_content)
                return {"text": text, "format": "word", "s3_bucket": s3_bucket, "s3_key": s3_key}

            elif file_extension in ['.txt', '.text']:
                text = file_content.decode('utf-8', errors='ignore')
                return {"text": text, "format": "text", "s3_bucket": s3_bucket, "s3_key": s3_key}

            else:
                return {
                    "error": f"Unsupported file format: {file_extension}. Supported: .pdf, .xlsx, .xls, .docx, .txt",
                    "text": "",
                    "format": "unknown"
                }

        except Exception as e:
            print(f"Error extracting text from S3: {e}")
            return {"error": str(e), "text": "", "format": "unknown"}

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file content

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text
        """
        import PyPDF2

        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())

            text = '\n'.join(text_parts)
            print(f"✓ Extracted {len(text)} characters from PDF ({len(pdf_reader.pages)} pages)")
            return text

        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return f"[PDF extraction error: {str(e)}]"

    def _extract_from_excel(self, file_content: bytes) -> str:
        """
        Extract text from Excel file content
        Converts tables and cells into structured text

        Args:
            file_content: Excel file bytes

        Returns:
            Extracted text in structured format
        """
        import pandas as pd

        try:
            excel_file = BytesIO(file_content)

            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')

            text_parts = []
            for sheet_name, df in excel_data.items():
                text_parts.append(f"\n{'='*60}")
                text_parts.append(f"SHEET: {sheet_name}")
                text_parts.append(f"{'='*60}\n")

                # Convert DataFrame to formatted text
                # Use markdown-like table format for better LLM understanding
                if not df.empty:
                    # Add column headers
                    headers = ' | '.join(str(col) for col in df.columns)
                    text_parts.append(headers)
                    text_parts.append('-' * len(headers))

                    # Add rows
                    for idx, row in df.iterrows():
                        row_text = ' | '.join(str(val) if pd.notna(val) else '' for val in row)
                        text_parts.append(row_text)

                    text_parts.append("")  # Empty line between sheets

            text = '\n'.join(text_parts)
            print(f"✓ Extracted {len(text)} characters from Excel ({len(excel_data)} sheets)")
            return text

        except Exception as e:
            print(f"Error extracting text from Excel: {e}")
            return f"[Excel extraction error: {str(e)}]"

    def _extract_from_word(self, file_content: bytes) -> str:
        """
        Extract text from Word (.docx) file content

        Args:
            file_content: Word file bytes

        Returns:
            Extracted text
        """
        from docx import Document

        try:
            word_file = BytesIO(file_content)
            doc = Document(word_file)

            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    text_parts.append(paragraph.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n{'='*60}")
                text_parts.append(f"TABLE {table_idx + 1}")
                text_parts.append(f"{'='*60}\n")

                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():  # Skip empty rows
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line after table

            text = '\n'.join(text_parts)
            print(f"✓ Extracted {len(text)} characters from Word document ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
            return text

        except Exception as e:
            print(f"Error extracting text from Word: {e}")
            return f"[Word extraction error: {str(e)}]"

    def extract_text_from_local(self, file_path: str) -> Dict[str, str]:
        """
        Extract text from a local file
        Useful for testing without S3

        Args:
            file_path: Local file path

        Returns:
            Dictionary with text, format, and optional error
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            with open(file_path, 'rb') as f:
                file_content = f.read()

            # Route to appropriate extractor
            if file_extension == '.pdf':
                text = self._extract_from_pdf(file_content)
                return {"text": text, "format": "pdf"}

            elif file_extension in ['.xlsx', '.xls']:
                text = self._extract_from_excel(file_content)
                return {"text": text, "format": "excel"}

            elif file_extension == '.docx':
                text = self._extract_from_word(file_content)
                return {"text": text, "format": "word"}

            elif file_extension in ['.txt', '.text']:
                text = file_content.decode('utf-8', errors='ignore')
                return {"text": text, "format": "text"}

            else:
                return {
                    "error": f"Unsupported file format: {file_extension}",
                    "text": "",
                    "format": "unknown"
                }

        except Exception as e:
            print(f"Error extracting text from local file: {e}")
            return {"error": str(e), "text": "", "format": "unknown"}
